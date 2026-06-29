import asyncio
import io
import logging
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any, Final
from urllib.parse import quote

from fastapi import APIRouter, Depends, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from langchain_chroma import Chroma
from starlette.requests import Request

from app.core.config import settings
from app.core.dependencies import get_vectorstore
from app.core.enums import DocumentFormat
from app.core.limiter import limiter
from app.exceptions import DocumentLoadError, UnsupportedFormatError, VectorStoreError
from app.models.schemas import PRDExportRequest, UploadResponse
from app.rag.vectorstore import add_document_to_vectorstore, clear_vectorstore
from app.security import authenticate
from app.security.auth import get_user_rate_limit
from app.security.rate_limiter import user_rate_limiter

logger = logging.getLogger(__name__)

business_audit_logger = logging.getLogger("audit.business")
business_audit_logger.setLevel(logging.INFO)
if not business_audit_logger.handlers:
    _audit_handler = logging.StreamHandler()
    _audit_handler.setFormatter(
        logging.Formatter(
            "[%(asctime)s] [BIZ_AUDIT] %(message)s",
            datefmt="%Y-%m-%d %H:%M:%S",
        )
    )
    business_audit_logger.addHandler(_audit_handler)

SUPPORTED_FORMATS: Final = {
    DocumentFormat.PDF,
    DocumentFormat.TXT,
    DocumentFormat.DOCX,
    DocumentFormat.MD,
}

__all__ = ["router"]

router = APIRouter()


def _markdown_to_docx(markdown_content: str, feature_name: str) -> io.BytesIO:
    from docx import Document

    doc = Document()
    doc.add_heading(feature_name or "文档", level=0)

    if not markdown_content or not markdown_content.strip():
        doc.add_paragraph("（无内容）")
    else:
        for line in markdown_content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@router.post("/docs/upload", response_model=UploadResponse, summary="上传文档")
@limiter.limit(f"{settings.RATE_LIMIT_RPM}/minute")
async def upload_document(
    request: Request,
    file: UploadFile,
    department: str = Form(default="通用", description="文档所属部门"),
    vectorstore: Chroma = Depends(get_vectorstore),
    user: dict = Depends(authenticate),
) -> UploadResponse:
    try:
        logger.info("收到文件上传：%s，部门：%s", file.filename, department)

        # 按用户限流
        if settings.ENABLE_RATE_LIMIT:
            user_limit = get_user_rate_limit(user)
            user_rate_limiter.check(key=user["hashed_key"], limit=user_limit)

        if file.filename is None:
            raise UnsupportedFormatError(
                format="未知",
                details="文件名不能为空",
            )

        file_ext = Path(file.filename).suffix.lower()

        match file_ext:
            case DocumentFormat.PDF | DocumentFormat.TXT | DocumentFormat.DOCX | DocumentFormat.MD:
                pass
            case _:
                raise UnsupportedFormatError(
                    format=file_ext,
                    details=(
                        f"仅支持 {DocumentFormat.PDF}、{DocumentFormat.TXT}、"
                        f"{DocumentFormat.DOCX} 和 {DocumentFormat.MD} 格式"
                    ),
                )

        knowledge_dir = Path(settings.KNOWLEDGE_DIR)
        knowledge_dir.mkdir(parents=True, exist_ok=True)

        file_path = knowledge_dir / file.filename
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        logger.info("文件已保存到：%s", file_path)

        import time

        start_time = time.monotonic()
        chunks_added = await asyncio.to_thread(
            add_document_to_vectorstore, vectorstore, file_path, department
        )
        processing_time_ms = (time.monotonic() - start_time) * 1000

        message = f"文档 {file.filename} 已成功上传并入库"
        logger.info("%s，添加 %d 个文本块，部门：%s", message, chunks_added, department)

        business_audit_logger.info(
            {
                "action": "upload_document",
                "filename": file.filename,
                "department": department,
                "chunks_added": chunks_added,
                "timestamp": datetime.now().isoformat(),
            }
        )

        return UploadResponse(
            filename=file.filename,
            chunks_added=chunks_added,
            message=message,
            department=department,
            processing_time_ms=round(processing_time_ms, 1),
            acceleration_mode="cloud_api",
        )

    except UnsupportedFormatError as exc:
        logger.error("格式错误：%s", exc)
        raise HTTPException(status_code=400, detail=str(exc))
    except DocumentLoadError as exc:
        logger.error("文档加载错误：%s", exc)
        raise HTTPException(status_code=400, detail=str(exc))
    except VectorStoreError as exc:
        logger.error("向量库错误：%s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    except Exception as exc:
        logger.error("上传错误：%s", exc)
        raise HTTPException(status_code=500, detail=f"文档上传失败：{exc!s}")


@router.delete("/docs/clear", summary="清空向量库")
@limiter.limit("5/minute")
async def clear_all_docs(
    request: Request,
    vectorstore: Chroma = Depends(get_vectorstore),
    user: dict = Depends(authenticate),
) -> dict[str, Any]:
    try:
        logger.info("收到清空向量库请求")
        await asyncio.to_thread(clear_vectorstore, vectorstore)

        business_audit_logger.info(
            {
                "action": "clear_vectorstore",
                "timestamp": datetime.now().isoformat(),
            }
        )

        return {"message": "向量库已清空"}
    except VectorStoreError as exc:
        logger.error("向量库错误：%s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    except Exception as exc:
        logger.error("清空向量库错误：%s", exc)
        raise HTTPException(status_code=500, detail=f"清空向量库失败：{exc!s}")


@router.get("/docs/count", summary="查询向量库文档数量")
async def get_doc_count(
    vectorstore: Chroma = Depends(get_vectorstore),
    user: dict = Depends(authenticate),
) -> dict[str, Any]:
    try:
        count = vectorstore._collection.count()
        logger.info("向量库文档数量：%d", count)
        return {"count": count}
    except VectorStoreError as exc:
        logger.error("向量库错误：%s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    except Exception as exc:
        logger.error("查询文档数量错误：%s", exc)
        raise HTTPException(status_code=500, detail=f"查询文档数量失败：{exc!s}")


@router.post("/docs/export/prd", summary="导出 PRD 文档为 Word")
@limiter.limit(f"{settings.RATE_LIMIT_RPM}/minute")
async def export_prd_document(
    request: Request,
    body: PRDExportRequest,
    user: dict = Depends(authenticate),
) -> StreamingResponse:
    try:
        logger.info("收到 PRD 导出请求，功能名称：%s", body.feature_name)

        if not body.content or not body.content.strip():
            raise HTTPException(status_code=400, detail="导出内容不能为空")

        docx_buffer = await asyncio.to_thread(_markdown_to_docx, body.content, body.feature_name)

        safe_filename = body.feature_name.replace(" ", "_")

        encoded_filename = quote(safe_filename + ".docx")

        logger.info("PRD 文档生成完成，文件名：%s.docx", safe_filename)

        business_audit_logger.info(
            {
                "action": "export_prd",
                "feature_name": body.feature_name,
                "timestamp": datetime.now().isoformat(),
            }
        )

        return StreamingResponse(
            docx_buffer,
            media_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            headers={
                "Content-Disposition": (
                    f"attachment; filename=\"export.docx\"; filename*=UTF-8''{encoded_filename}"
                ),
            },
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("PRD 导出失败：%s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=f"PRD 文档导出失败：{exc!s}")
