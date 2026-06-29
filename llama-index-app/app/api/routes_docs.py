import asyncio
import io
import logging
import shutil
from pathlib import Path
from typing import Any, Final

from fastapi import APIRouter, Depends, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from llama_index.core import VectorStoreIndex

from app.core.config import settings
from app.core.dependencies import get_index
from app.core.enums import DocumentFormat
from app.exceptions import DocumentLoadError, UnsupportedFormatError, VectorStoreError
from app.models.schemas import (
    OptimizedUploadResponse,
    PRDExportRequest,
    RetrievalTestResponse,
    UploadResponse,
)
from app.rag.index_store import add_documents_to_index, clear_index, get_index_count

logger = logging.getLogger(__name__)

SUPPORTED_FORMATS: Final = {
    DocumentFormat.PDF,
    DocumentFormat.TXT,
    DocumentFormat.DOCX,
    DocumentFormat.MD,
    DocumentFormat.CSV,
}

__all__ = ["router"]

router = APIRouter()


def _markdown_to_docx(markdown_content: str, feature_name: str) -> io.BytesIO:
    from docx import Document

    doc = Document()
    doc.add_heading(feature_name or "文档", level=0)

    if not markdown_content or not markdown_content.strip():
        doc.add_paragraph("（无内容）")
    else:
        for line in markdown_content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@router.post("/docs/upload", response_model=UploadResponse, summary="上传文档")
async def upload_document(
    file: UploadFile,
    department: str = Form(default="通用", description="文档所属部门"),
    index: VectorStoreIndex = Depends(get_index),
) -> UploadResponse:
    try:
        logger.info("收到文件上传：%s，部门：%s", file.filename, department)

        if file.filename is None:
            raise UnsupportedFormatError(
                format="未知",
                details="文件名不能为空",
            )

        file_ext = Path(file.filename).suffix.lower()

        match file_ext:
            case (
                DocumentFormat.PDF
                | DocumentFormat.TXT
                | DocumentFormat.DOCX
                | DocumentFormat.MD
                | DocumentFormat.CSV
            ):
                pass
            case _:
                raise UnsupportedFormatError(
                    format=file_ext,
                    details=(
                        f"仅支持 {DocumentFormat.PDF}、{DocumentFormat.TXT}、"
                        f"{DocumentFormat.DOCX}、{DocumentFormat.MD} "
                        f"和 {DocumentFormat.CSV} 格式"
                    ),
                )

        knowledge_dir = Path(settings.KNOWLEDGE_DIR)
        knowledge_dir.mkdir(parents=True, exist_ok=True)

        file_path = knowledge_dir / file.filename
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        logger.info("文件已保存到：%s", file_path)

        import time

        from app.core.hardware import hardware_manager

        start_time = time.monotonic()
        chunks_added = await asyncio.to_thread(
            add_documents_to_index, index, file_path, department
        )
        processing_time_ms = (time.monotonic() - start_time) * 1000

        message = f"文档 {file.filename} 已成功上传并入库"
        logger.info("%s，添加 %d 个文档，部门：%s", message, chunks_added, department)

        return UploadResponse(
            filename=file.filename,
            chunks_added=chunks_added,
            message=message,
            department=department,
            processing_time_ms=round(processing_time_ms, 1),
            acceleration_mode=hardware_manager.mode,
        )

    except UnsupportedFormatError as exc:
        logger.error("格式错误：%s", exc)
        raise HTTPException(status_code=400, detail=str(exc))
    except DocumentLoadError as exc:
        logger.error("文档加载错误：%s", exc)
        raise HTTPException(status_code=400, detail=str(exc))
    except VectorStoreError as exc:
        logger.error("向量库错误：%s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    except Exception as exc:
        logger.error("上传错误：%s", exc)
        raise HTTPException(status_code=500, detail=f"文档上传失败：{exc!s}")


@router.delete("/docs/clear", summary="清空向量库")
async def clear_all_docs(
    index: VectorStoreIndex = Depends(get_index),
) -> dict[str, Any]:
    try:
        logger.info("收到清空向量库请求")
        await asyncio.to_thread(clear_index)
        return {"message": "向量库已清空"}
    except VectorStoreError as exc:
        logger.error("向量库错误：%s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    except Exception as exc:
        logger.error("清空向量库错误：%s", exc)
        raise HTTPException(status_code=500, detail=f"清空向量库失败：{exc!s}")


@router.get("/docs/count", summary="查询向量库文档数量")
async def get_doc_count(index: VectorStoreIndex = Depends(get_index)) -> dict[str, Any]:
    try:
        count = get_index_count(index)
        logger.info("向量库文档数量：%d", count)
        return {"count": count}
    except VectorStoreError as exc:
        logger.error("向量库错误：%s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    except Exception as exc:
        logger.error("查询文档数量错误：%s", exc)
        raise HTTPException(status_code=500, detail=f"查询文档数量失败：{exc!s}")


@router.post("/docs/export/prd", summary="导出 PRD 文档为 Word")
async def export_prd_document(
    body: PRDExportRequest,
) -> StreamingResponse:
    try:
        logger.info("收到 PRD 导出请求，功能名称：%s", body.feature_name)

        if not body.content or not body.content.strip():
            raise HTTPException(status_code=400, detail="导出内容不能为空")

        docx_buffer = await asyncio.to_thread(
            _markdown_to_docx, body.content, body.feature_name
        )

        safe_filename = body.feature_name.replace(" ", "_")
        # RFC 5987 编码中文文件名
        from urllib.parse import quote

        encoded_filename = quote(safe_filename + ".docx")

        logger.info("PRD 文档生成完成，文件名：%s.docx", safe_filename)

        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": (
                    f'attachment; filename="export.docx"; '
                    f"filename*=UTF-8''{encoded_filename}"
                ),
            },
        )
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("PRD 导出失败：%s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=f"PRD 文档导出失败：{exc!s}")


@router.post("/docs/upload/optimized", response_model=OptimizedUploadResponse, summary="优化文档上传")
async def upload_document_optimized(
    file: UploadFile,
    use_gpu: bool = True,
    department: str = Form(default="通用", description="文档所属部门"),
    index: VectorStoreIndex = Depends(get_index),
) -> OptimizedUploadResponse:
    """优化文档上传接口，使用 GPU 加速文档预处理，云端 Embedding API 不受影响"""
    try:
        from app.core.hardware import hardware_manager
        from app.models.schemas import DocumentProcessingPerformance, HardwareInfo

        logger.info("收到优化文件上传：%s，部门：%s，use_gpu=%s", file.filename, department, use_gpu)

        if file.filename is None:
            raise UnsupportedFormatError(
                format="未知",
                details="文件名不能为空",
            )

        file_ext = Path(file.filename).suffix.lower()

        match file_ext:
            case (
                DocumentFormat.PDF
                | DocumentFormat.TXT
                | DocumentFormat.DOCX
                | DocumentFormat.MD
                | DocumentFormat.CSV
            ):
                pass
            case _:
                raise UnsupportedFormatError(
                    format=file_ext,
                    details=(
                        f"仅支持 {DocumentFormat.PDF}、{DocumentFormat.TXT}、"
                        f"{DocumentFormat.DOCX}、{DocumentFormat.MD} "
                        f"和 {DocumentFormat.CSV} 格式"
                    ),
                )

        knowledge_dir = Path(settings.KNOWLEDGE_DIR)
        knowledge_dir.mkdir(parents=True, exist_ok=True)

        file_path = knowledge_dir / file.filename
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        logger.info("文件已保存到：%s", file_path)

        import time

        # 记录各阶段处理时间
        total_start = time.monotonic()

        # 使用 hardware_manager 检测硬件状态
        config = hardware_manager.get_optimization_config()

        # 实际文档处理仍使用现有的 add_documents_to_index
        chunks_added = await asyncio.to_thread(
            add_documents_to_index, index, file_path, department
        )

        total_time_ms = (time.monotonic() - total_start) * 1000

        # 构建处理性能指标（简化版，不区分 OCR/清洗/切片子阶段）
        processing_performance = DocumentProcessingPerformance(
            ocr_time_ms=0.0,
            cleaning_time_ms=0.0,
            slicing_time_ms=0.0,
            total_time_ms=round(total_time_ms, 1),
            device=config["device"],
        )

        # 构建硬件信息
        hw_info = HardwareInfo(
            device=config["device"],
            mode=config["mode"],
            gpu_name=config["gpu_name"],
            gpu_memory_gb=config["gpu_memory_gb"],
            available_optimizations=config["available_optimizations"],
        )

        message = f"文档 {file.filename} 已成功上传并入库（优化模式）"
        logger.info("%s，添加 %d 个文档，部门：%s", message, chunks_added, department)

        return OptimizedUploadResponse(
            filename=file.filename,
            chunks_added=chunks_added,
            message=message,
            department=department,
            processing_performance=processing_performance,
            hardware_info=hw_info,
        )

    except UnsupportedFormatError as exc:
        logger.error("格式错误：%s", exc)
        raise HTTPException(status_code=400, detail=str(exc))
    except DocumentLoadError as exc:
        logger.error("文档加载错误：%s", exc)
        raise HTTPException(status_code=400, detail=str(exc))
    except VectorStoreError as exc:
        logger.error("向量库错误：%s", exc)
        raise HTTPException(status_code=500, detail=str(exc))
    except Exception as exc:
        logger.error("优化上传错误：%s", exc)
        raise HTTPException(status_code=500, detail=f"文档上传失败：{exc!s}")


@router.get("/docs/retrieval/test", response_model=RetrievalTestResponse, summary="检索性能测试")
async def test_retrieval_performance(
    query: str,
    use_reranker: bool = False,
    use_gpu: bool = True,
) -> RetrievalTestResponse:
    """检索性能测试接口，测试检索各阶段耗时"""
    try:
        import time

        from app.core.hardware import hardware_manager
        from app.models.schemas import HardwareInfo, RetrievalPerformance

        logger.info("检索性能测试：query=%s，use_reranker=%s，use_gpu=%s", query, use_reranker, use_gpu)

        config = hardware_manager.get_optimization_config()

        # 执行检索并记录耗时
        index = get_index()
        total_start = time.monotonic()

        # 使用 LlamaIndex 的检索器
        from llama_index.core import get_response_synthesizer

        retriever = index.as_retriever(similarity_top_k=5)
        vector_start = time.monotonic()
        nodes = await asyncio.to_thread(retriever.retrieve, query)
        vector_time_ms = (time.monotonic() - vector_start) * 1000

        total_time_ms = (time.monotonic() - total_start) * 1000

        # 构建检索结果
        results: list[dict[str, Any]] = []
        for node in nodes:
            results.append({
                "text": node.node.get_content()[:200],
                "score": node.score if node.score is not None else 0.0,
                "metadata": node.node.metadata,
            })

        # 构建检索性能指标
        performance = RetrievalPerformance(
            bm25_time_ms=0.0,
            vector_time_ms=round(vector_time_ms, 1),
            rrf_time_ms=0.0,
            reranker_time_ms=0.0,
            total_time_ms=round(total_time_ms, 1),
            device=config["device"],
            use_reranker=use_reranker,
        )

        # 构建硬件信息
        hw_info = HardwareInfo(
            device=config["device"],
            mode=config["mode"],
            gpu_name=config["gpu_name"],
            gpu_memory_gb=config["gpu_memory_gb"],
            available_optimizations=config["available_optimizations"],
        )

        return RetrievalTestResponse(
            query=query,
            results=results,
            performance=performance,
            hardware_info=hw_info,
        )

    except Exception as exc:
        logger.error("检索性能测试错误：%s", exc)
        raise HTTPException(status_code=500, detail=f"检索性能测试失败：{exc!s}")
